from io import BytesIO
from typing import List

from docx import Document
import fitz  # PyMuPDF

from app.services.interfaces.ITextExtractor import ITextExtractor
from app.services.interfaces.IOCR import IOCR


class TextExtractor(ITextExtractor):
    def __init__(self, ocr_service: IOCR):
        self._ocr_service = ocr_service

    def parse_docx(self, document_bytes: bytes) -> str:
        doc = Document(BytesIO(document_bytes))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text)
                parts.append("\t".join(row_text))

        return "\n".join(parts).strip()

    def parse_pdf(self, document_bytes: bytes) -> str:
        parts: List[str] = []
        with fitz.open(stream=document_bytes, filetype="pdf") as pdf:
            for page in pdf:
                page_text = page.get_text("text").strip()
                if page_text:
                    parts.append(page_text)

                for img in page.get_images(full=True):
                    xref = img[0]
                    try:
                        base_image = pdf.extract_image(xref)
                        image_bytes: bytes = base_image.get("image", b"")
                        if image_bytes:
                            _, ocr_text = self._ocr_service.extract_text(image_bytes)
                            if ocr_text and ocr_text.strip():
                                parts.append(ocr_text.strip())
                    except Exception:
                        continue

        return "\n".join([p for p in parts if p]).strip()